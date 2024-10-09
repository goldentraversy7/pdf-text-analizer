import csv
from bs4 import BeautifulSoup
from docx import Document
import re
import PyPDF2
import json
from datetime import datetime
from app.mongo import get_db
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from difflib import SequenceMatcher
from flask import session
from unidecode import unidecode
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk

WORD = re.compile(r"\w+")
# Descargar stopwords y tokenizer si no lo has hecho antesimport nltk
nltk.download('stopwords')
nltk.download('punkt')
nltk.download('punkt_tab')

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from collections import Counter
import math
import re

def get_pdf_text(pdf_path):
    text = ""
    with open(pdf_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def get_docx_text(file_path):
    doc = Document(file_path)

    # Extract text from the DOCX file
    full_content = []

    for para in doc.paragraphs:
        if para.text.strip():  # Check if paragraph is not empty
            # Check if the paragraph has any numbering
            if para.style.name.startswith('List'):
                if para.style.name == 'List Number':
                    numbering = 'Number: '  # Define a prefix for numbered lists
                elif para.style.name == 'List Bullet':
                    numbering = 'Bullet: '  # Define a prefix for bulleted lists
                else:
                    numbering = ''  # No numbering
                
                content = f"{numbering}{para.text}" 
            else:
                content = para.text
            full_content.append(content)
    
    return '\n'.join(full_content)

def clean_text(str):
    cite_pattern = r"【\d+†source】|【\d+:\d+†source】|【\d+:\d+†fuente】|【.*?】"
    return re.sub(cite_pattern, "", str)


def find_setencia_list(list):
    db = get_db()
    collection = db["sentencias"]

    json_data = []
    for item in list:
        doc = collection.find_one({"providencia": item['providencia']})
        if doc:
            json_data.append(
                {
                    "providencia": doc["providencia"],
                    # 'tipo': doc['tipo'],
                    # 'ano': doc['ano'],
                    "fecha_sentencia": doc["fecha_sentencia"],
                    # 'tema': doc['tema'],
                    "derechos": doc["derechos"],  # se cambio 'derechos' por 'derechos'
                    "magistrado": doc["magistrado"],
                    # 'fecha_publicada': doc['fecha_publicada'],
                    "expediente": doc["expediente"],
                    "url": doc["url"],
                }
            )
    return json_data


def html_to_text(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    return soup.get_text()


def add_html_to_docx(soup, doc):
    for element in soup:
        if element == "\n":
            continue
        if element.name == "p":
            paragraph = doc.add_paragraph(element.get_text())
            run = paragraph.runs[0]
            run.font.size = Pt(11.5)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h1":
            paragraph = doc.add_heading(level=1)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(20)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h2":
            paragraph = doc.add_heading(level=2)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(16)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h3":
            paragraph = doc.add_heading(level=3)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(14)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h4":
            paragraph = doc.add_heading(level=3)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(13)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "strong":
            run = doc.add_paragraph().add_run(element.get_text())
            paragraph_format = run.paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "em":
            run = doc.add_paragraph().add_run(element.get_text())
            run.italic = True
            run.font.size = Pt(11.5)
            paragraph_format = run.paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "ul":
            list_items = element.find_all("li")
            for i, li in enumerate(list_items):
                paragraph = doc.add_paragraph(li.get_text(), style="ListBullet")
                run = paragraph.runs[0]
                run.font.size = Pt(11.5)
                paragraph_format = paragraph.paragraph_format
                if i == 0:
                    paragraph_format.space_before = Pt(6)
                if i == len(list_items) - 1:
                    paragraph_format.space_after = Pt(6)
        elif element.name == "ol":
            list_items = element.find_all("li")
            for i, li in enumerate(list_items):
                paragraph = doc.add_paragraph(li.get_text(), style="ListNumber")
                run = paragraph.runs[0]
                run.font.size = Pt(11.5)
                paragraph_format = paragraph.paragraph_format
                if i == 0:
                    paragraph_format.space_before = Pt(6)
                if i == len(list_items) - 1:
                    paragraph_format.space_after = Pt(6)

        elif isinstance(element, str):
            paragraph = doc.add_paragraph(element)
            run = paragraph.runs[0]
            run.font.size = Pt(11.5)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after


def create_docx_from_html(html_content, filename):
    soup = BeautifulSoup(html_content, "lxml")

    doc = Document()
    add_html_to_docx(soup.body.contents, doc)

    doc.save(filename)


def get_constitution(str):
    pattern = r"Artículo (?P<number>\d+)"  # Named group to capture the number
    matches = re.findall(pattern, str)
    print("Found:", matches)

    return matches


def proccess_code(codigo):
    # Eliminar acentos y convertir a minúsculas
    codigo = unidecode(codigo.lower())

    codigo = re.sub(
        r"^\d+\.\s*", "", codigo
    ).strip()  # Eliminar números al inicio
    codigo = re.sub(
        r"[^A-Za-z0-9]", "-", codigo
    )  # Reemplazar caracteres no alfanuméricos por guiones
    codigo = re.sub(r"de-", "-", codigo)  # Eliminar 'de-' seguido de un guion
    codigo = re.sub(
        r"-+", " ", codigo
    )  # Reemplazar múltiples guiones consecutivos por un solo guion
    codigo = re.sub(
        r"\s+", " ", codigo
    )  # Unificar espacios múltiples
    codigo = re.sub(
        r"-(\d{2})(\d{2})$", r"-\2", codigo
    )  # Convertir años de 4 dígitos a 2 dígitos si es necesario
    return codigo


def parse_date(date_str):
    return datetime.strptime(date_str, "%Y-%m-%d")

# Función para calcular la similitud entre dos textos
def es_similar(texto1, texto2):
    return SequenceMatcher(None, texto1, texto2).ratio()

def es_quick_similar(texto1, texto2):
    def get_cosine(vec1, vec2):
        intersection = set(vec1.keys()) & set(vec2.keys())
        numerator = sum([vec1[x] * vec2[x] for x in intersection])

        sum1 = sum([vec1[x] ** 2 for x in list(vec1.keys())])
        sum2 = sum([vec2[x] ** 2 for x in list(vec2.keys())])
        denominator = math.sqrt(sum1) * math.sqrt(sum2)

        if not denominator:
            return 0.0
        else:
            return float(numerator) / denominator

    def text_to_vector(text):
        # Lista de stopwords en español
        stop_words = set(stopwords.words('spanish'))
        # Tokenizamos el texto y eliminamos stopwords y cualquier palabra no alfanumérica
        words = word_tokenize(text.lower(), language='spanish')
        words_filtered = [word for word in words if word.isalnum() and word not in stop_words]
        return Counter(words_filtered)

    def keyword_match(texto1, texto2):
        # Lista de palabras clave importantes
        palabras_clave = ["fertilización", "reproductiva", "consentimiento", "revocación", "madre", "Camila", "Andrés"]
        
        # Convertimos ambos textos a conjuntos de palabras
        palabras1 = set(word_tokenize(texto1.lower(), language='spanish'))
        palabras2 = set(word_tokenize(texto2.lower(), language='spanish'))
        
        # Buscamos coincidencias de palabras clave
        coincidencias = palabras1 & palabras2 & set(palabras_clave)
        
        # Devolvemos la proporción de coincidencias con respecto a las palabras clave
        return len(coincidencias) / len(palabras_clave) if palabras_clave else 0

    # Vectorizamos los textos usando el método text_to_vector
    vector1 = text_to_vector(texto1)
    vector2 = text_to_vector(texto2)

    # Calculamos la similitud del coseno entre los vectores
    cosine = get_cosine(vector1, vector2)

    # Calculamos la similitud basada en palabras clave
    keyword_similarity = keyword_match(texto1, texto2)

    # Combinamos ambas métricas: similitud del coseno y palabras clave
    final_similarity = (cosine + keyword_similarity) / 2

    return final_similarity


# Lista de verbos comunes para identificar peticiones (en infinitivo o imperativo)
VERBOS_COMUNES = r"\b(solicitar|realizar|suministrar|ordenar|iniciar|conceder|proteger|amparar|otorgar|dictaminar|evaluar|exigir|pedir|valorar|determinar|textr)\b"


def buscar_patrones_en_texto(TextSentTemp):
    print(
        f"TextSentTemp contiene {len(TextSentTemp)} elementos para análisis de sentencias adjuntas"
    )

    # Verificar que TextSentTemp sea una lista válida
    if not isinstance(TextSentTemp, list) or not TextSentTemp:
        print("Error: TextSentTemp no es una lista válida o está vacía.")
        return []

    # Inicializar la lista de sentencias coincidentes
    salida1 = []

    # Conectar a la base de datos MongoDB
    db = get_db()
    collection = db["sentencias"]
    
    # Diccionario para almacenar las providencias ya procesadas
    providencias_procesadas = {}

    # Crear un pipeline que busque coincidencias de las sentencias en la base de datos
    pipeline = [
        {"$project": {"_id": 1, "providencia": 1, "fecha_sentencia": 1}},
        {"$sort": {"fecha_sentencia": -1}},
    ]

    # Iteramos directamente sobre el cursor sin convertirlo en lista
    for doc in collection.aggregate(pipeline):
        providencia_db = doc["providencia"].lower().strip()
        providencia_db_normalizada = re.sub(
            r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", providencia_db
        )

        # Verificar si ya se procesó esta providencia
        if providencia_db_normalizada in providencias_procesadas:
            continue  # Saltar si ya fue procesada

        for sentencia_adjunta in TextSentTemp:
            sentencia_adjunta_normalizada = (
                re.sub(r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", sentencia_adjunta)
                .lower()
                .strip()
            )
            if sentencia_adjunta_normalizada == providencia_db_normalizada:
                print(
                    f"Coincidencia encontrada: '{providencia_db_normalizada}' en documento ID: {doc['_id']}'"
                )
                salida1.append(
                    {
                        "_id": doc["_id"],
                        "providencia": doc["providencia"],
                        "fecha_sentencia": doc["fecha_sentencia"],
                    }
                )
                # Marcar providencia como procesada inmediatamente después de encontrar la coincidencia
                providencias_procesadas[providencia_db_normalizada] = True
                break  # Salir del bucle interno si ya se encontró una coincidencia para esta providencia

    print(f"Total de coincidencias encontradas salida1: {len(salida1)}")
    return salida1


def normalize_pdf_resume(pdf_resume):
    sections = pdf_resume.split("\n")

    normalizadas = []
    for section in sections:
        normalizada = re.sub(
            r"^\b(primero|segundo|tercero|cuarto|quinto|sexto|séptimo|octavo|noveno|décimo||[a-zA-Z]\)|\d+[*,+]+[\.\)]|•|-)\b[\s]*",
            "",
            section,
            flags=re.IGNORECASE,
        ).strip()
        normalizada = re.sub(
            r"\b(Derechos Fundamentales Invocados|Hechos Relevantes|Peticiones|Pruebas Adjuntas|Sentencias)\b",
            "",
            normalizada,
            flags=re.IGNORECASE,
        ).strip()
        normalizada = re.sub(
            r"^\d+\.\s*", "", normalizada
        ).strip()  # Eliminar números al inicio
        normalizada = re.sub(
            r"[\W_]+$", "", normalizada
        ).strip()  # Eliminar caracteres especiales al final
        normalizada = re.sub(
            r"\s+", " ", normalizada
        ).strip()  # Unificar espacios múltiples

        if normalizada:
            normalizadas.append(normalizada)

    return normalizadas


def normalize_tema(tema):
    sections = re.split(r"[.]", tema)

    normalizadas = []
    for section in sections:
        normalizada = re.sub(
            r"^\b(primero|segundo|tercero|cuarto|quinto|sexto|séptimo|octavo|noveno|décimo||[a-zA-Z]\)|\d+[*,+]+[\.\)]|•|-)\b[\s]*",
            "",
            section,
            flags=re.IGNORECASE,
        ).strip()

        normalizada = re.sub(
            r"^\d+\.\s*", "", normalizada
        ).strip()  # Eliminar números al inicio
        normalizada = re.sub(
            r"[\W_]+$", "", normalizada
        ).strip()  # Eliminar caracteres especiales al final
        normalizada = re.sub(
            r"\s+", " ", normalizada
        ).strip()  # Unificar espacios múltiples

        if normalizada:
            normalizadas.append(normalizada)

    return normalizadas


def buscar_tema_en_mongo(pdf_resumen, limitar_busqueda=True):
    # Verificar si hay pdf_resumen fundamentales invocados
    if not pdf_resumen:
        print(
            "Error: No se encontraron los datos de pdf_resumen fundamentales invocados en la sesión."
        )
        return []

    # Normalizar pdf_resumen
    pdf_resumen_normalizados = proccess_code(pdf_resumen)

    # Conectar a la base de datos MongoDB
    db = get_db()
    collection = db["sentencias"]

    # Preparamos el pipeline de búsqueda
    pipeline = [
        {"$project": {"_id": 1, "providencia": 1, "fecha_sentencia": 1, "tema": 1}},
        {"$sort": {"fecha_sentencia": -1}},  # Ordenamos por la fecha de sentencia
    ]

    salida2 = []
    SentTemaTempFiles = []

    # Diccionario para almacenar el primer documento por providencia
    procesadas_providencias = {}

    for doc in collection.aggregate(pipeline):
        providencia = doc.get("providencia")

        # Si ya procesamos un documento con la misma providencia, lo saltamos
        if providencia in procesadas_providencias:
            continue

        tema_db = doc.get("tema", "").lower().strip()  # Convertimos tema a minúsculas
        tema_db_normalizados = proccess_code(tema_db)

        similitud = es_quick_similar(pdf_resumen_normalizados, tema_db_normalizados)

        salida2.append(
            {
                "_id": doc["_id"],
                "providencia": providencia,
                "fecha_sentencia": doc["fecha_sentencia"],
                "similitud": similitud,
            }
        )

        # Marcar la providencia como procesada inmediatamente después de encontrar la coincidencia
        procesadas_providencias[providencia] = True

    # Ordenamos la salida por similitud y limitamos a los 10 mejores resultados
    salida2 = sorted(salida2, key=lambda x: x["similitud"], reverse=True)[:10]

    return salida2



def buscar_derechos_en_mongo(salida2, TextDerecTemp):
    print(
        f"TextDerecTemp contiene {len(TextDerecTemp)} elementos como entrada análisis derechos"
    )

    if not TextDerecTemp:
        print(
            "Error: No se encontraron los datos de derechos fundamentales invocados en la sesión."
        )
        return salida2  # Si no hay derechos, retornamos `salida2` sin modificaciones

    # Accedemos a la base de datos
    db = get_db()
    collection = db["sentencias"]

    salida3 = []
    # coincidencias_derechos = 0

    # Recorrer los IDs obtenidos en salida2
    for doc_info in salida2:
        doc_id = doc_info["_id"]  # Asumimos que en salida2 cada documento tiene '_id'
        doc = collection.find_one({"_id": doc_id})
        if doc:
            derechos_db = (
                doc.get("derechos", "").lower().strip()
            )  # Convertimos derechos a minúsculas
            derechos_db_normalizado = re.sub(
                r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", derechos_db
            )  # Normalizamos el formato

            for derecho in TextDerecTemp:
                derecho_normalizado = re.sub(
                    r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", derecho.lower().strip()
                )

                # Verificamos la similitud entre el derecho en la base de datos y el derecho de la lista de entrada
                similitud = es_similar(derecho_normalizado, derechos_db_normalizado)

                if similitud > 0.2:
                    # coincidencias_derechos += 1
                    print(
                        f"Similitud alta entre '{derecho_normalizado}' y '{derechos_db_normalizado}' (Similitud: {similitud})"
                    )

                    # Si se encuentra una coincidencia, agregamos el documento a la salida
                    salida3.append(
                        {
                            "_id": doc["_id"],
                            "providencia": doc["providencia"],
                            "fecha_sentencia": doc["fecha_sentencia"],
                            "similitud": similitud,
                        }
                    )

    # Ordenamos por similitud y limitamos a 10 resultados
    salida3 = sorted(salida3, key=lambda x: x["similitud"], reverse=True)[:10]

    return salida3


def buscar_peticiones_con_ids(salida3, TextPetTemp):
    print(
        f"TextPetTemp contiene {len(TextPetTemp)} elementos como entrada análisis ids"
    )

    if not TextPetTemp:
        print(
            "Error: No se encontraron los datos de peticiones fundamentales invocados en la sesión."
        )
        return salida3  # Si no hay peticiones, retornamos `salida3` sin modificaciones

    # Conectar a la base de datos MongoDB
    db = get_db()
    collection = db["sentencias"]

    salida4 = salida3.copy()  # Creating a copy

    SentPetTempFiles = []

    # Recorrer los IDs obtenidos en salida3
    for doc_info in salida3:
        doc_id = doc_info["_id"]  # Asumimos que en salida3 cada documento tiene '_id'
        doc = collection.find_one({"_id": doc_id})
        if doc:
            texto = doc.get(
                "texto", ""
            ).lower()  # Texto completo del documento en minúsculas
            coincidencias_peticiones = 0

            print(f"Analizando documento con ID {doc_id}")

            # Comparar las peticiones con el texto del documento
            for peticion in TextPetTemp:
                peticion_normalizada = peticion.lower().strip()

                # Coincidencia directa
                if peticion_normalizada in texto:
                    coincidencias_peticiones += 1
                    print(
                        f"Coincidencia directa para petición: '{peticion_normalizada}' en documento ID: {doc_id}"
                    )

                # Coincidencia parcial utilizando similitud
                oraciones_documento = re.split(
                    r"[.!?]", texto
                )  # Dividir el documento en oraciones
                for oracion in oraciones_documento:
                    similitud = es_similar(peticion_normalizada, oracion.strip())
                    if similitud > 0.65:  # Umbral de similitud flexible
                        coincidencias_peticiones += 1
                        print(
                            f"Similitud entre '{peticion}' y '{oracion[:50]}...' en documento ID: {doc_id}, Similitud: {similitud}"
                        )

            if coincidencias_peticiones > 0:
                SentPetTempFiles.append((doc, coincidencias_peticiones))

    if SentPetTempFiles:
        # Ordenamos los resultados por el número de coincidencias y limitamos el número de resultados
        SentPetTempFiles = sorted(SentPetTempFiles, key=lambda x: x[1], reverse=True)[
            :7
        ]
        salida4 = [
            {
                "_id": doc["_id"],
                "providencia": doc["providencia"],
                "fecha_sentencia": doc["fecha_sentencia"],
            }
            for doc, _ in SentPetTempFiles
        ]

    return salida4


def get_sentencia(TextResumenTemp, TextDerecTemp, TextPetTemp, TextSentTemp):
    print("Iniciando get_sentencia...")
    print(f"TextDerecTemp: {TextDerecTemp}")
    print(f"TextPetTemp: {TextPetTemp}")
    print(f"TextSentTemp: {TextSentTemp}")

    # Parte 1: Búsqueda de patrones en el texto (pattern matching)
    salida1 = buscar_patrones_en_texto(TextSentTemp)
    print(f"Resultados de buscar_patrones_en_texto: {salida1}")

    # Parte 2: Búsqueda de tema en la base de datos (pdf_resume)
    salida2 = buscar_tema_en_mongo(TextResumenTemp)
    print(f"Resultados de buscar_tema_en_mongo: Salida2: {salida2}")

    # Parte 3: Búsqueda de derechos en la base de datos (TextDerecTemp)
    salida3 = buscar_derechos_en_mongo(salida2, TextDerecTemp)
    print(f"Resultados de buscar_derechos_en_mongo: Salida3: {salida3}")

    # Parte 4: Búsqueda en Mongo de los IDs obtenidos y matching con TextPetTemp
    salida4 = buscar_peticiones_con_ids(salida3, TextPetTemp)
    print(f"Resultados de buscar_peticiones_con_ids: {salida4}")

    print(
        f"Resultados de get_sentencia: \n salida1 => {salida1} \n salida4 => {salida4}"
    )

    # Combinar las salidas sin duplicar elementos
    sentencia_list = []
    for salida in [salida1, salida4]:
        for item in salida:
            if item not in sentencia_list:
                sentencia_list.append(item)

    # Imprimir el formato pipeline para verificar la salida
    print(f"Formato final para el frontend (pipeline): {sentencia_list}")

    # Crear el formato esperado por el frontend basado en el pipeline
    db = get_db()
    collection = db["sentencias"]

    # Crear un pipeline para devolver los documentos que coincidan con sentencia_list
    pipeline = [
        {
            "$match": {
                "_id": {
                    "$in": [item["_id"] for item in sentencia_list if "_id" in item]
                }
            }
        },
        {"$project": {"_id": 1, "providencia": 1, "fecha_sentencia": 1}},
        {"$sort": {"fecha_sentencia": -1}},
    ]

    # Ejecutar el pipeline y devolver los resultados en el formato esperado
    sentencia_list = list(collection.aggregate(pipeline=pipeline))

    sentencia_list = [
        {
            "providencia": doc["providencia"],
            "fecha_sentencia": doc["fecha_sentencia"],
        }
        for doc in sentencia_list
    ]

    return sentencia_list


def generate_evidence_checklist(text):
    match = re.search(r"```json(.*?)```", text, re.DOTALL)
    if match:
        json_text = match.group(1).strip()
    else:
        print("No JSON block found.")
        return
    # Cargar el texto JSON en un objeto
    json_object = json.loads(json_text)

    # Asegurar que cada evidencia solo tenga dos checkboxes
    for item in json_object:
        if "evidencias" in item:
            for evidencia in item["evidencias"]:
                # Añadir las opciones "Cumple" y "No cumple" por cada evidencia
                evidencia["checkboxes"] = ["Cumple", "No cumple"]

    return json_object


def get_history(user, title=""):
    db = get_db()
    collection = db["saves"]
    pipeline = []
    if title == "":
        pipeline.append({"$project": {"_id": 1, "user": 1, "modifiedAt": 1}})
        pipeline.append({"$match": {"user": user}})
        pipeline.append({"$sort": {"modifiedAt": -1}})
        pipeline.append({"$limit": 1})
    else:
        pipeline.append(
            {"$project": {"_id": 1, "user": 1, "modifiedAt": 1, "title": 1}}
        )
        pipeline.append({"$match": {"user": user, "title": title}})
    find_data = list(collection.aggregate(pipeline=pipeline))
    if len(find_data) > 1:
        return "Error"
    elif len(find_data) == 1:
        history = collection.find_one({"_id": find_data[0]["_id"]})
        print(history)
        return history
    else:
        return "No data"


def get_current_state(user):
    db = get_db()
    collection = db["current_state"]
    history = collection.find_one({"user": user})
    if history:
        history.pop("_id", None)
        return history
    else:
        collection.insert_one({"user": user})
        return {"user": user}


def update_current_state(user, field, data):
    db = get_db()
    collection = db["current_state"]
    query_filter = {"user": user}
    update_date = {}
    update_date[field] = data
    update_operation = {}
    update_operation["$set"] = update_date
    collection.update_one(query_filter, update_operation)


def get_current_data_field(user, field):
    db = get_db()
    collection = db["current_state"]
    history = collection.find_one({"user": user})
    if field in history:
        return history[field]
    else:
        return ""


def get_settings(field):
    db = get_db()
    collection = db["settings"]
    settings = collection.find_one()
    return settings[field]


def get_title_list(user):
    db = get_db()
    collection = db["results"]
    pipeline = [{"$project": {"title": 1, "user": 1}}, {"$match": {"user": user}}]
    res = list(collection.aggregate(pipeline=pipeline))
    return_data = []
    for each in res:
        return_data.append(each["title"])
    return return_data


def save_text(user, title):
    loading_data = get_current_state(user)
    loading_data["title"] = title
    loading_data["modifiedAt"] = datetime.now()

    # Asegúrate de que las evidencias se guarden correctamente
    evidencias_cumplen = session.get("evidencias_cumplen", [])
    evidencias_no_cumplen = session.get("evidencias_no_cumplen", [])

    loading_data["evidencias_cumplen"] = evidencias_cumplen
    loading_data["evidencias_no_cumplen"] = evidencias_no_cumplen

    db = get_db()
    results = db["results"]
    currents = db["current_state"]

    # Verifica si ya existe un análisis con el mismo título
    pipeline = [
        {"$project": {"_id": 1, "title": 1, "user": 1}},
        {"$match": {"user": user, "title": title}},
    ]
    res = list(results.aggregate(pipeline=pipeline))

    # Si ya existe, lo eliminamos para sobreescribir
    if len(res) > 0:
        results.delete_one({"user": user, "title": title})

    # Insertar los datos en `results`
    results.insert_one(loading_data)

    # También actualizamos el `current_state`
    currents.delete_one({"user": user})
    currents.insert_one(loading_data)

    return {"message": "sucess"}, 200


def set_text(user, title):
    db = get_db()
    results = db["results"]
    currents = db["current_state"]

    # Buscar los datos guardados en `results`
    set_data = results.find_one({"user": user, "title": title})

    # Si los datos se encuentran, los transferimos a `current_state`
    if set_data:
        currents.delete_one({"user": user})
        currents.insert_one(set_data)

        # Asegúrate de que las evidencias también se carguen en la sesión
        session["evidencias_cumplen"] = set_data.get("evidencias_cumplen", [])
        session["evidencias_no_cumplen"] = set_data.get("evidencias_no_cumplen", [])

        return {"message": "sucess"}, 200
    else:
        return {"message": "Análisis no encontrado"}, 404


def reset_current_state(user):
    db = get_db()
    currents = db["current_state"]
    currents.delete_one({"user": user})  # Esto debe borrar el estado anterior

    return
